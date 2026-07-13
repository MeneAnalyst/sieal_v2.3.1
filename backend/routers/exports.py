"""
Multi-format export router: CSV, XLSX, TXT, DOCX, PDF, JSON.

Section 7 of the instruction set requires a single unified endpoint
(`GET /api/export/{dataset}?format=...`) in addition to the legacy
per-format paths the original vanilla-JS frontend already links to.
Both are kept so neither UI breaks.
"""
import csv, io, json
from datetime import date, timedelta
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, HTMLResponse, JSONResponse
from sqlalchemy.orm import Session
from sqlalchemy import func

from database import get_db
from models import ARTClient, Batch, DispenseRecord, Drug, VLResult, ExpiryLoss, StockTransfer
from routers.auth import current_user

router = APIRouter()

# ── helpers ──────────────────────────────────────────────────────
def patient_rows(db, fid):
    q = db.query(ARTClient).filter(ARTClient.is_active == 1)
    if fid: q = q.filter(ARTClient.facility_id == fid)
    return q.all()

def stock_rows(db):
    today = date.today()
    return db.query(Batch).filter(Batch.quantity_remaining > 0).order_by(Batch.expiry_date).all()

def dispense_rows(db):
    return db.query(DispenseRecord).order_by(DispenseRecord.dispense_date.desc()).limit(5000).all()

PAT_HEADERS = ["ART Number","TB Number","Full Name","DOB","Gender","Combination",
               "Visit Type","Initiation Date","Progress Status","CD4","VL Result",
               "VL Suppressed","Adherence","Stock Status","Last Visit","Next Appointment","ECI"]

def pat_row(c):
    return [c.art_number, c.tb_number or "", c.full_name,
            c.date_of_birth or "", c.gender or "",
            c.treatment_combination or "", c.visit_type,
            c.initiation_date or "", c.progress_status,
            c.cd4_count or "", c.vl_result or "",
            "Yes" if c.vl_suppressed else "No",
            f"{c.adherence_score}%" if c.adherence_score else "",
            c.stock_status or "IN", c.last_visit or "", c.next_appointment or "",
            "Yes" if c.is_eci_flag else "No"]

STOCK_HEADERS = ["Drug","Batch Number","Expiry Date","Days to Expiry",
                 "Qty Received","Qty Remaining","Supplier","Alert Status"]

def stk_row(b, today):
    days = (b.expiry_date - today).days
    status = "RED" if days <= 30 else "AMBER" if days <= 90 else "GREEN"
    return [b.drug.name if b.drug else "", b.batch_number,
            b.expiry_date, days, b.quantity_received, b.quantity_remaining,
            b.supplier or "", status]

DISP_HEADERS = ["Date","Client","ART Number","Drug","Batch","Qty","Dispensed By"]

def disp_row(r):
    return [r.dispense_date,
            r.client.full_name if r.client else "",
            r.client.art_number if r.client else "",
            r.batch.drug.name if r.batch and r.batch.drug else "",
            r.batch.batch_number if r.batch else "",
            r.quantity, r.dispensed_by or ""]

def loss_rows(db):
    return db.query(ExpiryLoss).order_by(ExpiryLoss.loss_date.desc()).limit(5000).all()

LOSS_HEADERS = ["Date", "Drug", "Batch", "Qty Lost", "Reason", "Notes"]

def loss_row(l):
    return [l.loss_date,
            l.batch.drug.name if l.batch and l.batch.drug else "",
            l.batch.batch_number if l.batch else "",
            l.quantity_lost, l.reason_code or "", l.notes or ""]

def eci_rows(db, fid):
    q = db.query(ARTClient).filter(ARTClient.is_active == 1, ARTClient.is_eci_flag == 1)
    if fid: q = q.filter(ARTClient.facility_id == fid)
    return q.order_by(ARTClient.eci_flagged_date.desc()).all()

ECI_HEADERS = ["ART Number", "Full Name", "Status", "CD4", "VL Result", "Reason", "Flagged Date"]

def eci_row(c):
    return [c.art_number, c.full_name, c.progress_status, c.cd4_count or "",
            c.vl_result or "", c.eci_reason or "", c.eci_flagged_date or ""]

def transfer_rows(db):
    return db.query(StockTransfer).order_by(StockTransfer.request_date.desc()).limit(2000).all()

TRANSFER_HEADERS = ["Date", "Drug", "Donor", "Receiver", "Requested", "Approved", "Status", "Repaid"]

def transfer_row(t):
    return [t.request_date, t.drug.name if t.drug else "",
            t.donor.name if t.donor else "", t.receiver.name if t.receiver else "",
            t.quantity_requested, t.quantity_approved or "", t.status, t.quantity_repaid or 0]

# ── CSV ───────────────────────────────────────────────────────────
@router.get("/csv/patients")
def csv_patients(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    buf = io.StringIO(); w = csv.writer(buf); w.writerow(PAT_HEADERS)
    [w.writerow(pat_row(c)) for c in patient_rows(db, session.get("facility_id"))]
    buf.seek(0)
    return StreamingResponse(buf, media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=patients.csv"})

@router.get("/csv/stock")
def csv_stock(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    today = date.today(); buf = io.StringIO(); w = csv.writer(buf); w.writerow(STOCK_HEADERS)
    [w.writerow(stk_row(b, today)) for b in stock_rows(db)]
    buf.seek(0)
    return StreamingResponse(buf, media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=stock.csv"})

@router.get("/csv/dispenses")
def csv_dispenses(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    buf = io.StringIO(); w = csv.writer(buf); w.writerow(DISP_HEADERS)
    [w.writerow(disp_row(r)) for r in dispense_rows(db)]
    buf.seek(0)
    return StreamingResponse(buf, media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=dispenses.csv"})

# ── TXT ───────────────────────────────────────────────────────────
@router.get("/txt/patients")
def txt_patients(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    lines = ["SIEAL — Patient Register", f"Generated: {date.today()}", "="*80, ""]
    for c in patient_rows(db, session.get("facility_id")):
        lines += [f"ART: {c.art_number}  Name: {c.full_name}  Status: {c.progress_status}",
                  f"  Combo: {c.treatment_combination or '—'}  CD4: {c.cd4_count or '—'}  VL: {c.vl_result or '—'}",
                  f"  Next Appt: {c.next_appointment or '—'}  ECI: {'YES' if c.is_eci_flag else 'No'}", ""]
    buf = io.BytesIO("\n".join(lines).encode())
    return StreamingResponse(buf, media_type="text/plain",
        headers={"Content-Disposition": "attachment; filename=patients.txt"})

# ── XLSX ──────────────────────────────────────────────────────────
@router.get("/xlsx/patients")
def xlsx_patients(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        wb = Workbook(); ws = wb.active; ws.title = "Patients"
        teal = "0D9488"; hdr_fill = PatternFill("solid", fgColor=teal)
        ws.append(PAT_HEADERS)
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF", name="Arial")
            cell.fill = hdr_fill; cell.alignment = Alignment(horizontal="center")
        for c in patient_rows(db, session.get("facility_id")):
            ws.append([str(x) for x in pat_row(c)])
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = 18
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=patients.xlsx"})
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

@router.get("/xlsx/stock")
def xlsx_stock(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        today = date.today(); wb = Workbook(); ws = wb.active; ws.title = "Stock"
        teal = "0D9488"; hdr_fill = PatternFill("solid", fgColor=teal)
        ws.append(STOCK_HEADERS)
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF", name="Arial")
            cell.fill = hdr_fill
        for b in stock_rows(db):
            row = stk_row(b, today)
            ws.append([str(x) for x in row])
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = 20
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=stock.xlsx"})
    except ImportError:
        return {"error": "openpyxl not installed"}

# ── DOCX ──────────────────────────────────────────────────────────
@router.get("/docx/patients")
def docx_patients(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading("SIEAL — Patient Register", 0)
        doc.add_paragraph(f"Generated: {date.today()} | Facility: {session.get('facility_name','')}")
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["ART Number","Name","Combination","CD4","VL","Status","Next Appt"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
        for c in patient_rows(db, session.get("facility_id")):
            row = table.add_row().cells
            row[0].text = c.art_number or ""
            row[1].text = c.full_name or ""
            row[2].text = c.treatment_combination or ""
            row[3].text = str(c.cd4_count or "")
            row[4].text = str(c.vl_result or "")
            row[5].text = c.progress_status or ""
            row[6].text = str(c.next_appointment or "")
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=patients.docx"})
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

# ── HTML PRINT (PDF via browser) ──────────────────────────────────
@router.get("/pdf/patients", response_class=HTMLResponse)
def pdf_patients(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    patients = patient_rows(db, session.get("facility_id"))
    rows = "".join(f"""<tr>
        <td>{c.art_number}</td><td>{c.full_name}</td>
        <td>{c.treatment_combination or '—'}</td>
        <td style="color:{'red' if c.cd4_count and c.cd4_count < 200 else 'black'}">{c.cd4_count or '—'}</td>
        <td style="color:{'red' if not c.vl_suppressed and c.vl_result else 'green'}">{c.vl_result or '—'}</td>
        <td>{c.progress_status}</td>
        <td>{'YES' if c.is_eci_flag else '—'}</td>
        <td>{c.next_appointment or '—'}</td>
    </tr>""" for c in patients)
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"/>
    <title>Patient Register</title>
    <style>body{{font-family:Arial;font-size:11px;margin:20px}}
    h1{{color:#0D9488;font-size:18px}}
    table{{border-collapse:collapse;width:100%}}
    th{{background:#0D9488;color:#fff;padding:6px;text-align:left}}
    td{{border:1px solid #ddd;padding:5px}}
    tr:nth-child(even){{background:#f0fdfa}}
    @media print{{button{{display:none}}}}</style>
    </head><body>
    <button onclick="window.print()" style="background:#0D9488;color:#fff;border:none;padding:8px 18px;border-radius:6px;cursor:pointer;margin-bottom:14px">Print / Save PDF</button>
    <h1>SIEAL — Patient Register</h1>
    <p>Generated: {date.today()} | {session.get('facility_name','')} | {len(patients)} patients</p>
    <table><thead><tr>
    <th>ART Number</th><th>Name</th><th>Combination</th><th>CD4</th><th>VL</th><th>Status</th><th>ECI</th><th>Next Appt</th>
    </tr></thead><tbody>{rows}</tbody></table></body></html>"""
    return html

# ── JSON ──────────────────────────────────────────────────────────
@router.get("/json/summary")
def json_summary(db: Session = Depends(get_db), session: dict = Depends(current_user)):
    fid = session.get("facility_id"); today = date.today()
    q = db.query(ARTClient).filter(ARTClient.is_active == 1)
    if fid: q = q.filter(ARTClient.facility_id == fid)
    total = q.count(); active = q.filter(ARTClient.progress_status == "ACTIVE").count()
    ltfu = q.filter(ARTClient.progress_status == "LTFU").count()
    eci = q.filter(ARTClient.is_eci_flag == 1).count()
    supp = q.filter(ARTClient.vl_suppressed == 1).count()
    unsupp = q.filter(ARTClient.vl_suppressed == 0, ARTClient.vl_result != None).count()
    return {
        "generated": today.isoformat(), "facility": session.get("facility_name"),
        "patients": {"total": total, "active": active, "ltfu": ltfu, "eci_flagged": eci},
        "viral_load": {"suppressed": supp, "unsuppressed": unsupp,
                       "suppression_pct": round(supp/(supp+unsupp)*100,1) if supp+unsupp else 0},
    }


# ══════════════════════════════════════════════════════════════════
# Section 7 — Unified export endpoint: GET /api/export/{dataset}?format=
# dataset  ∈ {patients, stock, dispenses}
# format   ∈ {csv, xlsx, pdf, docx, txt, json}
# ══════════════════════════════════════════════════════════════════

DATASETS = {
    "patients":  {"headers": PAT_HEADERS,      "rows": lambda db, fid: [pat_row(c) for c in patient_rows(db, fid)],  "title": "Patient Register"},
    "stock":     {"headers": STOCK_HEADERS,    "rows": lambda db, fid: [stk_row(b, date.today()) for b in stock_rows(db)], "title": "Stock Inventory (FEFO)"},
    "dispenses": {"headers": DISP_HEADERS,     "rows": lambda db, fid: [disp_row(r) for r in dispense_rows(db)],     "title": "Dispense History"},
    "losses":    {"headers": LOSS_HEADERS,     "rows": lambda db, fid: [loss_row(l) for l in loss_rows(db)],         "title": "Expiry Loss Register"},
    "eci":       {"headers": ECI_HEADERS,      "rows": lambda db, fid: [eci_row(c) for c in eci_rows(db, fid)],      "title": "Early Case Identification Register"},
    "transfers": {"headers": TRANSFER_HEADERS, "rows": lambda db, fid: [transfer_row(t) for t in transfer_rows(db)], "title": "Network Transfer Register"},
}


def _facility_header_lines(session: dict) -> list[str]:
    """Common header block: facility + timestamp, reused by PDF/DOCX/TXT."""
    return [
        f"Facility: {session.get('facility_name') or '—'}",
        f"Generated: {date.today().isoformat()} {'':>0}",
    ]


def _reportlab_pdf(dataset_key: str, headers: list[str], rows: list, session: dict) -> io.BytesIO:
    """
    Section 7 — real PDF via reportlab (not the browser print-dialog HTML
    fallback below), with a facility header and generation timestamp as required.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), topMargin=14 * mm, bottomMargin=14 * mm,
                             leftMargin=12 * mm, rightMargin=12 * mm)
    styles = getSampleStyleSheet()
    title_style = styles["Title"]; title_style.textColor = colors.HexColor("#0F766E")
    meta_style = styles["Normal"]

    elements = [
        Paragraph(f"SIEAL — {DATASETS[dataset_key]['title']}", title_style),
        Paragraph(f"Facility: {session.get('facility_name') or '—'} &nbsp;|&nbsp; "
                  f"Generated: {date.today().isoformat()} &nbsp;|&nbsp; {len(rows)} records", meta_style),
        Spacer(1, 10),
    ]

    table_data = [headers] + [[str(v) if v is not None else "" for v in row] for row in rows]
    tbl = Table(table_data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0D9488")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E2E8F0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0FDFA")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(tbl)
    doc.build(elements)
    buf.seek(0)
    return buf


def _multisheet_xlsx(dataset_key: str, headers: list[str], rows: list, session: dict) -> io.BytesIO:
    """Section 7 — multi-sheet workbook: a 'Data' sheet plus a 'Summary' cover sheet."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    summary = wb.active
    summary.title = "Summary"
    teal_fill = PatternFill("solid", fgColor="0D9488")
    summary["A1"] = f"SIEAL — {DATASETS[dataset_key]['title']}"
    summary["A1"].font = Font(bold=True, size=14, color="0F766E")
    summary["A3"] = "Facility"; summary["B3"] = session.get("facility_name") or "—"
    summary["A4"] = "Generated"; summary["B4"] = date.today().isoformat()
    summary["A5"] = "Record count"; summary["B5"] = len(rows)
    for r in (3, 4, 5):
        summary[f"A{r}"].font = Font(bold=True)
    summary.column_dimensions["A"].width = 16
    summary.column_dimensions["B"].width = 32

    data_ws = wb.create_sheet("Data")
    data_ws.append(headers)
    for cell in data_ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = teal_fill
        cell.alignment = Alignment(horizontal="center")
    for row in rows:
        data_ws.append([str(v) if v is not None else "" for v in row])
    for col in data_ws.columns:
        data_ws.column_dimensions[col[0].column_letter].width = 18

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf


def _docx_export(dataset_key: str, headers: list[str], rows: list, session: dict) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(f"SIEAL — {DATASETS[dataset_key]['title']}", level=0)
    doc.add_paragraph(f"Facility: {session.get('facility_name') or '—'}  |  Generated: {date.today().isoformat()}  |  {len(rows)} records")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        if hdr_cells[i].paragraphs[0].runs:
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v) if v is not None else ""
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


@router.get("/{dataset}")
def unified_export(
    dataset: str,
    format: str = "csv",
    db: Session = Depends(get_db),
    session: dict = Depends(current_user),
):
    """
    Section 7: GET /api/export/{dataset}?format={csv|xlsx|pdf|docx|txt|json}
    dataset ∈ patients | stock | dispenses
    """
    if dataset not in DATASETS:
        raise HTTPException(404, f"Unknown dataset '{dataset}'. Choose from: {', '.join(DATASETS)}")
    meta = DATASETS[dataset]
    fid = session.get("facility_id")
    headers = meta["headers"]
    rows = meta["rows"](db, fid)
    fmt = format.lower()

    if fmt == "csv":
        buf = io.StringIO(); w = csv.writer(buf); w.writerow(headers)
        [w.writerow(r) for r in rows]
        buf.seek(0)
        return StreamingResponse(buf, media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={dataset}.csv"})

    if fmt == "json":
        payload = {
            "dataset": dataset, "generated": date.today().isoformat(),
            "facility": session.get("facility_name"), "count": len(rows),
            "headers": headers,
            "rows": [[str(v) if v is not None else None for v in r] for r in rows],
        }
        return JSONResponse(payload)

    if fmt == "txt":
        lines = [f"SIEAL — {meta['title']}", *_facility_header_lines(session), "=" * 80, ""]
        for r in rows:
            lines.append(" | ".join(f"{h}: {v}" for h, v in zip(headers, r)))
        buf = io.BytesIO("\n".join(lines).encode())
        return StreamingResponse(buf, media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename={dataset}.txt"})

    if fmt == "xlsx":
        try:
            buf = _multisheet_xlsx(dataset, headers, rows, session)
        except ImportError:
            raise HTTPException(500, "openpyxl not installed. Run: pip install openpyxl")
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={dataset}.xlsx"})

    if fmt == "docx":
        try:
            buf = _docx_export(dataset, headers, rows, session)
        except ImportError:
            raise HTTPException(500, "python-docx not installed. Run: pip install python-docx")
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={dataset}.docx"})

    if fmt == "pdf":
        try:
            buf = _reportlab_pdf(dataset, headers, rows, session)
        except ImportError:
            raise HTTPException(500, "reportlab not installed. Run: pip install reportlab")
        return StreamingResponse(buf, media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={dataset}.pdf"})

    raise HTTPException(400, f"Unsupported format '{format}'. Choose from: csv, xlsx, pdf, docx, txt, json")
